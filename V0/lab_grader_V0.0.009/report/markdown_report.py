"""
채점 결과 -> 보고서(Markdown 기본, python-docx 설치돼 있으면 Word도 옵션 생성).
정기점검 프로젝트에서 설계했던 "표준형/요약형/경영보고형" 개념 중 표준형 하나를 우선 구현.
"""
import datetime


def build_markdown_report(project_name, scored, ai_result=None):
    lines = []
    lines.append(f"# {project_name} 점검 보고서")
    lines.append(f"작성일: {datetime.date.today().isoformat()}")
    lines.append("")

    if ai_result:
        lines.append("## Executive Summary")
        lines.append(ai_result.get("summary", ""))
        lines.append(f"(분석 출처: {ai_result.get('source', 'unknown')})")
        lines.append("")

    lines.append("## 단계별 결과")
    lines.append("| 단계 | 상태 | PASS | TOTAL |")
    lines.append("|---|---|---|---|")
    for s in scored:
        lines.append(f"| {s['label']} | {s['status']} | {s['pass']} | {s['total']} |")
    lines.append("")

    lines.append("## 특이사항 상세")
    for s in scored:
        fails = [r for r in s.get("results", []) if r["result"] in ("FAIL", "UNKNOWN")]
        if not fails:
            continue
        lines.append(f"### {s['label']}")
        for r in fails:
            lines.append(f"- **{r['check']}** ({r.get('device','-')})  기대값: {r.get('expected')} / 실제: {r.get('actual')}")
        lines.append("")

    if ai_result and ai_result.get("top_priority"):
        lines.append("## 조치 권고 (우선순위 순)")
        for a in ai_result["top_priority"]:
            lines.append(f"{a['priority']}. **{a['device']} / {a['check']}** — {a.get('suggested_action', '')}")
        lines.append("")

    return "\n".join(lines)


def save_markdown_report(project_name, scored, ai_result, output_path):
    content = build_markdown_report(project_name, scored, ai_result)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    return output_path


def save_docx_report(project_name, scored, ai_result, output_path):
    """python-docx가 설치돼 있을 때만 동작. 없으면 안내만 하고 markdown으로 대체."""
    try:
        from docx import Document
    except ImportError:
        print("[안내] python-docx 미설치 — Word 보고서는 건너뛰고 Markdown만 생성됨(pip install python-docx로 추가 가능)")
        return None

    doc = Document()
    doc.add_heading(f"{project_name} 점검 보고서", level=1)
    doc.add_paragraph(f"작성일: {datetime.date.today().isoformat()}")

    if ai_result:
        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(ai_result.get("summary", ""))

    doc.add_heading("단계별 결과", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "단계", "상태", "PASS", "TOTAL"
    for s in scored:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = s["label"], s["status"], str(s["pass"]), str(s["total"])

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    import sys, os
    sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

    sample_scored = [
        {"label": "VLAN", "status": "COMPLETE", "pass": 18, "total": 18, "results": []},
        {"label": "STP", "status": "IN_PROGRESS", "pass": 3, "total": 14, "results": [
            {"stage": "STP", "device": "Core1", "check": "root_priority_vlan1_core1", "result": "FAIL", "expected": 4096, "actual": 32768},
        ]},
    ]
    from ai_analysis import rule_based
    ai_result = rule_based.analyze(sample_scored)
    md = build_markdown_report("LAB1 Campus", sample_scored, ai_result)
    print(md)
